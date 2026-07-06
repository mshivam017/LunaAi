import os
from pypdf import PdfReader
from docx import Document

class FileParserService:
    @classmethod
    def parse_file(cls, file_path: str) -> str:
        if not os.path.exists(file_path):
            return "Error: File not found."

        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext in [".txt", ".md", ".json", ".py", ".js", ".ts", ".html", ".css", ".csv"]:
                return cls._read_text_file(file_path)
            elif ext == ".pdf":
                return cls._read_pdf_file(file_path)
            elif ext in [".docx", ".doc"]:
                return cls._read_docx_file(file_path)
            elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
                return cls._read_image_mock_ocr(file_path)
            else:
                return f"Unsupported file type '{ext}'. Displaying binary preview placeholder."
        except Exception as e:
            return f"Error parsing file: {str(e)}"

    @staticmethod
    def _read_text_file(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            # Read first 10,000 characters to prevent memory blow-up on huge files
            content = f.read(10000)
            if len(content) >= 10000:
                content += "\n\n... [Content truncated due to size limits] ..."
            return content

    @staticmethod
    def _read_pdf_file(file_path: str) -> str:
        reader = PdfReader(file_path)
        text_content = []
        
        # Read max 3 pages for high speed and safe memory allocation on CPU
        max_pages = min(len(reader.pages), 3)
        for i in range(max_pages):
            page_text = reader.pages[i].extract_text()
            if page_text:
                text_content.append(f"--- Page {i+1} ---\n{page_text}")
                
        if len(reader.pages) > 3:
            text_content.append(f"\n... [Document truncated. Total pages: {len(reader.pages)}] ...")
            
        full_text = "\n".join(text_content).strip()
        # Strictly truncate full text to 3,000 characters max
        if len(full_text) > 3000:
            full_text = full_text[:3000] + "\n... [Content truncated] ..."
        return full_text if full_text else "No extractable text found in PDF."

    @staticmethod
    def _read_docx_file(file_path: str) -> str:
        doc = Document(file_path)
        text_content = []
        for para in doc.paragraphs[:100]: # Read max 100 paragraphs
            if para.text.strip():
                text_content.append(para.text)
                
        if len(doc.paragraphs) > 100:
            text_content.append("\n... [Document truncated due to size limits] ...")
            
        full_text = "\n".join(text_content).strip()
        return full_text if full_text else "No extractable text found in Word document."

    @staticmethod
    def _read_image_mock_ocr(file_path: str) -> str:
        filename = os.path.basename(file_path)
        file_size_kb = round(os.path.getsize(file_path) / 1024, 1)
        
        # Simple simulated intelligent OCR for testing and demos
        mock_ocr_descriptions = [
            f"Image file: {filename} ({file_size_kb} KB).\n\n"
            "[Intelligent Mock OCR active]\n"
            "Analyzed visual elements:\n"
            "- Core: User interface dashboard structure detected.\n"
            "- Recognized Text snippet: 'Luna AI Assistant - Settings Panel'.\n"
            "- Visual Theme: Dark-mode theme with vibrant blue and purple neon gradients.\n"
            "- Content: Contains interactive buttons, card components, and scrollable logs list.",
            
            f"Image file: {filename} ({file_size_kb} KB).\n\n"
            "[Intelligent Mock OCR active]\n"
            "Analyzed visual elements:\n"
            "- Core: Text-heavy invoice/document layout detected.\n"
            "- Recognized Text snippet: 'Invoice Summary - Total Due: $150.00'.\n"
            "- Layout: Column table layout tracking product IDs and order completion statuses.",
            
            f"Image file: {filename} ({file_size_kb} KB).\n\n"
            "[Intelligent Mock OCR active]\n"
            "Analyzed visual elements:\n"
            "- Core: Landscape screenshot or graphic art.\n"
            "- Main elements: High contrast abstract background with glowing lighting effects."
        ]
        
        # Select a mock response based on filename hash to keep it deterministic for the same file
        index = hash(filename) % len(mock_ocr_descriptions)
        return mock_ocr_descriptions[index]
